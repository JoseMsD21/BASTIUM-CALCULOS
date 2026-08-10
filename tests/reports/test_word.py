from docx import Document
from docx.shared import RGBColor

from app.reports.header import build_encabezado
from app.reports.word import WordReportGenerator


def _table_data():
    return [
        {
            "fecha": "15/04/2026",
            "concepto": "Abono a capital",
            "base_capital": "$1,500,000.50",
            "tasa": "1.50%",
            "interes": "$10,000.00",
            "indexacion": "$0.00",
            "pago": "$500,000.00",
            "saldo_capital": "$1,500,000.50",
            "saldo_interes": "$125,000.00",
            "saldo_total": "$1,625,000.50",
        }
    ]


def _summary():
    return {
        "total_abonos": "$500,000.00",
        "total_intereses_generados": "$10,000.00",
        "saldo_final_capital": "$1,500,000.50",
        "saldo_final_intereses": "$125,000.00",
        "saldo_final_indexacion": "$0.00",
        "gran_total_adeudado": "$1,625,000.50",
    }


def test_generate_incluye_titulo_encabezado_y_tabla_cronologica(tmp_path):
    ruta = tmp_path / "liquidacion.docx"
    generador = WordReportGenerator(str(ruta))
    encabezado = build_encabezado("2026-030", "Ana", "Luis", "Juzgado 5 Civil del Circuito")

    generador.generate(
        "LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", _summary(), _table_data(), encabezado
    )

    assert ruta.exists()
    documento = Document(str(ruta))
    texto_completo = "\n".join(parrafo.text for parrafo in documento.paragraphs)

    assert "LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA" in texto_completo
    assert "Radicado: 2026-030" in texto_completo
    assert "Ana vs. Luis" in texto_completo
    assert "Juzgado: Juzgado 5 Civil del Circuito" in texto_completo

    tabla_cronologia = documento.tables[1]
    encabezados_columna = [celda.text for celda in tabla_cronologia.rows[0].cells]
    assert encabezados_columna == [
        "Fecha",
        "Concepto",
        "Base Capital",
        "Tasa",
        "Interés",
        "Indexación/Sanciones",
        "Pago",
        "Saldo Capital",
        "Saldo Interés",
        "Saldo Total",
    ]
    fila_datos = [celda.text for celda in tabla_cronologia.rows[1].cells]
    assert fila_datos[4] == "$10,000.00"


def test_generate_sin_encabezado_no_falla(tmp_path):
    ruta = tmp_path / "liquidacion_sin_encabezado.docx"
    generador = WordReportGenerator(str(ruta))

    generador.generate(
        "LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", _summary(), _table_data()
    )

    assert ruta.exists()
    assert ruta.stat().st_size > 0


def test_generate_incluye_bloque_de_renta_liquida_cuando_se_provee(tmp_path):
    ruta = tmp_path / "liquidacion_renta.docx"
    generador = WordReportGenerator(str(ruta))
    renta_liquida = {
        "ingresos_netos": "$100,000,000.00",
        "renta_bruta": "$60,000,000.00",
        "renta_liquida": "$40,000,000.00",
        "hubo_perdida_liquida": "No",
        "renta_liquida_gravable": "$35,000,000.00",
    }

    generador.generate(
        "LIQUIDACIÓN DE OBLIGACIONES — ÁREA TRIBUTARIO",
        _summary(),
        _table_data(),
        renta_liquida=renta_liquida,
    )

    documento = Document(str(ruta))
    texto_completo = "\n".join(p.text for p in documento.paragraphs)
    assert "Depuración de Renta Líquida Gravable" in texto_completo
    tabla_renta_liquida = documento.tables[2]
    fila_gravable = [celda.text for celda in tabla_renta_liquida.rows[-1].cells]
    assert "$35,000,000.00" in fila_gravable


def test_generate_sin_renta_liquida_no_agrega_el_bloque(tmp_path):
    ruta = tmp_path / "liquidacion_sin_renta.docx"
    generador = WordReportGenerator(str(ruta))

    generador.generate(
        "LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", _summary(), _table_data()
    )

    documento = Document(str(ruta))
    assert len(documento.tables) == 2


def test_generate_resalta_en_rojo_las_filas_prescritas(tmp_path):
    # Sprint 42: ReportTableBuilder.build_matrix ya expone "prescrita" por fila
    # (ver tests/reports/test_table_builder.py) -- generate() debe resaltarla en
    # el Word (indicador visual, aqui: texto en rojo) sin excluirla de la tabla
    # ni tocar ningun monto.
    ruta = tmp_path / "liquidacion_prescrita.docx"
    generador = WordReportGenerator(str(ruta))
    datos = _table_data()
    datos[0]["prescrita"] = True

    generador.generate("LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", _summary(), datos)

    documento = Document(str(ruta))
    tabla_cronologia = documento.tables[1]
    fila_prescrita = tabla_cronologia.rows[1]

    # El texto de cada celda se conserva intacto (nada se excluye ni se pierde).
    assert fila_prescrita.cells[1].text == "Abono a capital"

    run_concepto = fila_prescrita.cells[1].paragraphs[0].runs[0]
    assert run_concepto.font.color.rgb == RGBColor(0xC0, 0x00, 0x00)


def test_generate_no_resalta_filas_sin_marcar(tmp_path):
    ruta = tmp_path / "liquidacion_sin_marca.docx"
    generador = WordReportGenerator(str(ruta))

    generador.generate(
        "LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", _summary(), _table_data()
    )

    documento = Document(str(ruta))
    tabla_cronologia = documento.tables[1]
    fila_normal = tabla_cronologia.rows[1]
    run_concepto = fila_normal.cells[1].paragraphs[0].runs[0]
    assert run_concepto.font.color.rgb is None


def test_generate_incluye_columna_de_indexacion_sanciones(tmp_path):
    ruta = tmp_path / "liquidacion.docx"
    generador = WordReportGenerator(str(ruta))

    generador.generate("LIQUIDACIÓN DE OBLIGACIONES — ÁREA TRIBUTARIO", _summary(), _table_data())

    documento = Document(str(ruta))
    tabla_cronologia = documento.tables[1]
    encabezados = [celda.text for celda in tabla_cronologia.rows[0].cells]
    assert "Indexación/Sanciones" in encabezados
    fila_datos = [celda.text for celda in tabla_cronologia.rows[1].cells]
    assert "$0.00" in fila_datos


def test_generate_muestra_saldo_a_favor_en_el_resumen_cuando_esta_presente(tmp_path):
    # Sprint 46: cuando ReportSummaryBuilder.build_summary agrega la clave
    # "saldo_a_favor" (solo si hay sobrepago), el Word debe mostrar una linea
    # "Saldo a favor del deudor" en la tabla de resumen ejecutivo.
    ruta = tmp_path / "liquidacion_saldo_a_favor.docx"
    generador = WordReportGenerator(str(ruta))
    summary = _summary()
    summary["saldo_a_favor"] = "$3,000,000.00"

    generador.generate(
        "LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", summary, _table_data()
    )

    documento = Document(str(ruta))
    tabla_resumen = documento.tables[0]
    filas = [(fila.cells[0].text, fila.cells[1].text) for fila in tabla_resumen.rows]
    assert ("Saldo a favor del deudor", "$3,000,000.00") in filas


def test_generate_no_muestra_saldo_a_favor_en_el_resumen_cuando_no_hay_sobrepago(tmp_path):
    ruta = tmp_path / "liquidacion_sin_saldo_a_favor.docx"
    generador = WordReportGenerator(str(ruta))

    generador.generate(
        "LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", _summary(), _table_data()
    )

    documento = Document(str(ruta))
    tabla_resumen = documento.tables[0]
    etiquetas = [fila.cells[0].text for fila in tabla_resumen.rows]
    assert "Saldo a favor del deudor" not in etiquetas


def test_generate_muestra_columna_de_saldo_a_favor_en_la_cronologia_cuando_hay_sobrepago(tmp_path):
    # ReportTableBuilder.build_matrix (Sprint 46) agrega "saldo_a_favor" a cada
    # fila -- el Word agrega la columna solo cuando el resumen indica que hubo
    # sobrepago, para no romper el layout de las liquidaciones normales.
    ruta = tmp_path / "liquidacion_cronologia_saldo_a_favor.docx"
    generador = WordReportGenerator(str(ruta))
    summary = _summary()
    summary["saldo_a_favor"] = "$3,000,000.00"
    datos = _table_data()
    datos[0]["saldo_a_favor"] = "$3,000,000.00"

    generador.generate("LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", summary, datos)

    documento = Document(str(ruta))
    tabla_cronologia = documento.tables[1]
    encabezados = [celda.text for celda in tabla_cronologia.rows[0].cells]
    assert "Saldo a favor" in encabezados
    indice_columna = encabezados.index("Saldo a favor")
    fila_datos = [celda.text for celda in tabla_cronologia.rows[1].cells]
    assert fila_datos[indice_columna] == "$3,000,000.00"


def test_generate_no_agrega_columna_de_saldo_a_favor_sin_sobrepago(tmp_path):
    ruta = tmp_path / "liquidacion_cronologia_sin_saldo_a_favor.docx"
    generador = WordReportGenerator(str(ruta))

    generador.generate(
        "LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", _summary(), _table_data()
    )

    documento = Document(str(ruta))
    tabla_cronologia = documento.tables[1]
    encabezados = [celda.text for celda in tabla_cronologia.rows[0].cells]
    assert "Saldo a favor" not in encabezados
    assert len(encabezados) == 10
