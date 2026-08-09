from datetime import date
from decimal import Decimal

from docx import Document
from reportlab.platypus import Table

from app.engine.financial.rate import Rate
from app.engine.liquidation.engine import LiquidationCore
from app.engine.reports.summary import ReportSummaryBuilder
from app.engine.reports.table_builder import ReportTableBuilder
from app.engine.temporal.schedulers.base import Event
from app.reports.pdf import JudicialPDFGenerator
from app.reports.word import WordReportGenerator
from app.views.liquidaciones import ResultadoLiquidacionView


def _liquidar_expediente_con_sobrepago():
    # Mismo escenario del bug real (auditoria 2026-07-21, ver
    # tests/liquidation/test_engine.py::test_engine_sobrepago_expone_remanente_como_saldo_a_favor):
    # un pago de $10.000.000 contra una deuda de $7.000.000 -- el excedente de
    # $3.000.000 debe quedar registrado como saldo a favor, no desaparecer.
    events = [
        Event(date=date(2026, 1, 1), payload={"amount": Decimal("7000000.00")}, event_type="INSTALLMENT"),
        Event(date=date(2026, 1, 10), payload={"amount": Decimal("10000000.00")}, event_type="PAYMENT"),
    ]
    control_rate = Rate.from_percent(Decimal("0.0"))
    engine = LiquidationCore(default_daily_rate=control_rate)
    return engine.process(events, cutoff_date=date(2026, 1, 10))


def test_liquidacion_con_sobrepago_expone_saldo_a_favor_en_el_resultado():
    resultado = _liquidar_expediente_con_sobrepago()

    assert resultado.total_saldo_a_favor() == Decimal("3000000.00")
    # El total de abonos aplicados NO debe incluir el excedente -- solo lo que
    # efectivamente se imputo a la deuda (ningun total existente cambia de valor).
    assert resultado.total_payments_applied() == Decimal("7000000.00")
    assert resultado.final_balance().total() == Decimal("0.00")


def test_liquidacion_con_sobrepago_aparece_en_el_pdf(tmp_path, monkeypatch):
    resultado = _liquidar_expediente_con_sobrepago()
    summary = ReportSummaryBuilder().build_summary(resultado)
    table_data = ReportTableBuilder().build_matrix(resultado)

    tabla_real = Table
    llamadas = []

    def _tabla_capturada(datos, *args, **kwargs):
        llamadas.append(datos)
        return tabla_real(datos, *args, **kwargs)

    monkeypatch.setattr("app.reports.pdf.Table", _tabla_capturada)

    ruta = tmp_path / "liquidacion_sobrepago.pdf"
    JudicialPDFGenerator(str(ruta)).generate(
        "LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", summary, table_data
    )

    assert ruta.exists()
    datos_resumen = llamadas[0]
    assert ["Saldo a favor del deudor", "$3,000,000.00"] in datos_resumen

    datos_cronologia = llamadas[1]
    indice_columna = datos_cronologia[0].index("Saldo a favor")
    fila_pago = next(
        fila for fila in datos_cronologia[1:] if fila[indice_columna] == "$3,000,000.00"
    )
    assert fila_pago is not None
    # El total adeudado sigue en cero -- el sobrepago no altera ningun otro monto.
    assert summary["gran_total_adeudado"] == "$0.00"


def test_liquidacion_con_sobrepago_aparece_en_el_word(tmp_path):
    resultado = _liquidar_expediente_con_sobrepago()
    summary = ReportSummaryBuilder().build_summary(resultado)
    table_data = ReportTableBuilder().build_matrix(resultado)

    ruta = tmp_path / "liquidacion_sobrepago.docx"
    WordReportGenerator(str(ruta)).generate(
        "LIQUIDACIÓN DE OBLIGACIONES — ÁREA CIVIL / FAMILIA", summary, table_data
    )

    documento = Document(str(ruta))
    tabla_resumen = documento.tables[0]
    filas_resumen = [(fila.cells[0].text, fila.cells[1].text) for fila in tabla_resumen.rows]
    assert ("Saldo a favor del deudor", "$3,000,000.00") in filas_resumen

    tabla_cronologia = documento.tables[1]
    encabezados = [celda.text for celda in tabla_cronologia.rows[0].cells]
    indice_columna = encabezados.index("Saldo a favor")
    valores_columna = [
        fila.cells[indice_columna].text for fila in tabla_cronologia.rows[1:]
    ]
    assert "$3,000,000.00" in valores_columna


def test_liquidacion_con_sobrepago_aparece_en_la_pantalla_de_resultado(qtbot):
    resultado = _liquidar_expediente_con_sobrepago()

    view = ResultadoLiquidacionView()
    qtbot.addWidget(view)
    view.show()
    view.mostrar(resultado, expediente_id=1)

    assert view.etiqueta_saldo_a_favor.isVisible()
    assert "3000000.00" in view.etiqueta_saldo_a_favor.text()
    # El saldo final mostrado sigue siendo el correcto (cero), sin contaminarse
    # con el saldo a favor.
    assert "Saldo final: 0.00" in view.etiqueta_saldo_final.text()
