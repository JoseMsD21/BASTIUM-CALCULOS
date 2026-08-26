from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Sprint 108 (pedido del despacho, 2026-08-22): mismo criterio de identidad
# visual que app/reports/pdf.py -- encabezado de columnas en extrabold crema
# sobre fondo negro, cuerpo crema con texto negro, fila de totales en
# borgoña negrita. "Ancizar Sans ExtraBold" es el nombre de familia real
# dentro de app/assets/fonts/AncizarSans-ExtraBold.ttf (ver metadata `name`
# del TTF) -- a diferencia de un PDF, un .docx no embebe la fuente, asi que
# Word la resuelve por nombre si esta instalada y si no cae a su sustituto
# por defecto, sin romper el documento.
_NOMBRE_FUENTE_EXTRABOLD = "Ancizar Sans ExtraBold"


def _sombrear_celda(celda, color_hex: str) -> None:
    sombreado = OxmlElement("w:shd")
    sombreado.set(qn("w:val"), "clear")
    sombreado.set(qn("w:fill"), color_hex)
    celda._tc.get_or_add_tcPr().append(sombreado)


def _escribir_celda(
    celda,
    texto: str,
    *,
    color: RGBColor | None = None,
    bold: bool = False,
    font_name: str | None = None,
) -> None:
    # Nota: no usar `celda.text = ""` antes de este `add_run` -- el setter de
    # `_Cell.text` en python-docx siempre agrega su propio run vacio primero
    # (`tc.clear_content()` + `add_p()` + `add_r()`), lo que dejaba el texto
    # con estilo como el *segundo* run de la celda en vez del primero. Las
    # celdas que llegan aca ya estan recien creadas por `add_table`/`add_row`
    # (un parrafo vacio sin runs), asi que no hace falta limpiar nada.
    run = celda.paragraphs[0].add_run(texto)
    if color is not None:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    if font_name is not None:
        run.font.name = font_name

# Peso relativo de cada columna de la tabla de cronologia -- mismo criterio y
# mismos pesos que _PESO_COLUMNA_CRONOLOGIA en app/reports/pdf.py (mantener
# ambos en sync si se ajusta uno). "Concepto" pesa mas por ser la unica columna
# con texto largo.
_PESO_COLUMNA_CRONOLOGIA = {
    "Fecha": 0.7,
    "Concepto": 2.0,
    "Base Capital": 0.95,
    "Tasa": 0.55,
    "Interés": 0.85,
    "Indexación/Sanciones": 1.05,
    "Pago": 0.85,
    "Saldo Capital": 0.95,
    "Saldo Interés": 0.9,
    "Saldo Total": 0.95,
    "Saldo a favor": 0.85,
}


class WordReportGenerator:
    def __init__(self, output_path: str):
        self.output_path = output_path
        self.c_burgundy = RGBColor(0xAE, 0x1C, 0x21)
        self.c_prescrita = RGBColor(0xC0, 0x00, 0x00)
        self.c_advertencia = RGBColor(0xED, 0x6C, 0x02)
        self.c_cream = RGBColor(0xF5, 0xF1, 0xE9)
        self.c_black_hex = "000000"
        self.c_cream_hex = "F5F1E9"

    def _anchos_columnas_cronologia(
        self, encabezados: list[str], ancho_disponible: Cm
    ) -> list[Cm]:
        pesos = [_PESO_COLUMNA_CRONOLOGIA.get(encabezado, 0.85) for encabezado in encabezados]
        total_pesos = sum(pesos)
        return [Cm((peso / total_pesos) * ancho_disponible.cm) for peso in pesos]

    def _fijar_ancho_columna(self, tabla, indice: int, ancho: Cm) -> None:
        # Quirk conocido de python-docx: asignar solo `table.columns[i].width` no
        # es suficiente para que Word respete el ancho al abrir el documento --
        # hay que fijarlo tambien celda por celda de esa columna.
        tabla.columns[indice].width = ancho
        for fila in tabla.rows:
            fila.cells[indice].width = ancho

    def generate(
        self,
        title: str,
        summary: dict,
        table_data: list,
        encabezado: dict | None = None,
        renta_liquida: dict | None = None,
        diferencia_recalculo: dict | None = None,
        cuerpo_legal: str | None = None,
        alertas: list[str] | None = None,
    ) -> None:
        documento = Document()

        # Horizontal (Sprint 50, hallazgo de la prueba practica Civil/Familia): la
        # tabla de cronologia tiene 10-11 columnas -- en vertical (Letter portrait,
        # ~15.9cm utiles) no cabe; en horizontal (~24.9cm utiles con margenes de
        # 1.5cm) si. Ademas de girar la pagina, hay que fijar el ancho de cada
        # columna explicitamente (ver _fijar_ancho_columna) porque el estilo
        # "Table Grid" por defecto usa autofit-to-contents, que ignora los margenes
        # de la pagina cuando el contenido es ancho -- esa era la causa real del
        # desborde, no la orientacion en si.
        seccion = documento.sections[0]
        seccion.orientation = WD_ORIENT.LANDSCAPE
        # python-docx no intercambia el ancho/alto de pagina automaticamente al
        # cambiar `orientation` -- hay que hacerlo a mano.
        seccion.page_width, seccion.page_height = seccion.page_height, seccion.page_width
        seccion.left_margin = Cm(1.5)
        seccion.right_margin = Cm(1.5)
        seccion.top_margin = Cm(1.5)
        seccion.bottom_margin = Cm(1.5)
        ancho_disponible = Cm(
            seccion.page_width.cm - seccion.left_margin.cm - seccion.right_margin.cm
        )

        parrafo_titulo = documento.add_paragraph()
        parrafo_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = parrafo_titulo.add_run(title)
        run_titulo.bold = True
        run_titulo.font.size = Pt(16)
        run_titulo.font.color.rgb = self.c_burgundy

        if encabezado:
            if encabezado.get("radicado"):
                documento.add_paragraph(f"Radicado: {encabezado['radicado']}")
            if encabezado.get("partes"):
                documento.add_paragraph(encabezado["partes"])
            if encabezado.get("juzgado"):
                documento.add_paragraph(f"Juzgado: {encabezado['juzgado']}")

        documento.add_paragraph()

        # Sprint 47: parrafo de fundamento legal para los memoriales de
        # recalculo historico (Art. 53 C.P. / Art. 151 CPACA) -- mismo patron
        # aditivo que diferencia_recalculo/renta_liquida.
        if cuerpo_legal:
            documento.add_paragraph(cuerpo_legal)
            documento.add_paragraph()

        # Sprint 77: LiquidationResult.alertas (Sprint 43) -- advertencias no
        # bloqueantes ("Doble Actualización Prohibida", "Techo de usura
        # alcanzado") que ya se muestran en pantalla (banner en
        # ResultadoLiquidacionView) pero no llegaban al PDF/Word exportado.
        # Mismo patron aditivo que cuerpo_legal/renta_liquida: solo aparece
        # cuando el llamador provee una lista no vacia.
        if alertas:
            parrafo_advertencias = documento.add_paragraph()
            run_advertencias = parrafo_advertencias.add_run("Advertencias")
            run_advertencias.bold = True
            run_advertencias.font.color.rgb = self.c_advertencia
            for alerta in alertas:
                parrafo_alerta = documento.add_paragraph()
                run_alerta = parrafo_alerta.add_run(f"⚠ {alerta}")
                run_alerta.font.color.rgb = self.c_advertencia
            documento.add_paragraph()

        filas_resumen = [
            ("Total Abonos Aplicados", summary["total_abonos"]),
            ("Intereses Generados", summary["total_intereses_generados"]),
            ("Saldo Final Capital", summary["saldo_final_capital"]),
            ("Saldo Final Intereses", summary["saldo_final_intereses"]),
            ("Saldo Final Indexación/Sanciones", summary["saldo_final_indexacion"]),
            ("GRAN TOTAL ADEUDADO", summary["gran_total_adeudado"]),
        ]
        # Sprint 46: el saldo a favor de un sobrepago (Sprint 23) solo se muestra
        # cuando ReportSummaryBuilder.build_summary lo agrego al diccionario (es
        # decir, cuando efectivamente hubo un sobrepago).
        if "saldo_a_favor" in summary:
            filas_resumen.append(("Saldo a favor del deudor", summary["saldo_a_favor"]))
        tabla_resumen = documento.add_table(rows=1, cols=2)
        tabla_resumen.style = "Table Grid"
        celdas_encabezado = tabla_resumen.rows[0].cells
        for celda, texto in zip(
            celdas_encabezado, ("Rubro Financiero", "Monto Liquidado"), strict=True
        ):
            _escribir_celda(
                celda, texto, color=self.c_cream, bold=True, font_name=_NOMBRE_FUENTE_EXTRABOLD
            )
            _sombrear_celda(celda, self.c_black_hex)
        indice_ultima_fila = len(filas_resumen) - 1
        for indice, (etiqueta, valor) in enumerate(filas_resumen):
            celdas_fila = tabla_resumen.add_row().cells
            es_fila_total = indice == indice_ultima_fila
            for celda, texto in zip(celdas_fila, (etiqueta, valor), strict=True):
                _escribir_celda(
                    celda,
                    texto,
                    color=self.c_burgundy if es_fila_total else None,
                    bold=es_fila_total,
                )
                _sombrear_celda(celda, self.c_cream_hex)

        documento.add_paragraph()
        parrafo_subtitulo = documento.add_paragraph()
        run_subtitulo = parrafo_subtitulo.add_run("Cronología Detallada de Imputaciones y Saldos")
        run_subtitulo.bold = True
        run_subtitulo.font.color.rgb = self.c_burgundy

        # Sprint 46: la columna "Saldo a favor" solo se agrega cuando el resumen
        # ya indico que hubo un sobrepago (mismo criterio que la linea del
        # resumen ejecutivo) -- asi el layout de las 6 areas no cambia para la
        # inmensa mayoria de liquidaciones, que nunca tienen sobrepago.
        mostrar_saldo_a_favor = "saldo_a_favor" in summary

        columnas_cronologia = [
            "Fecha",
            "Concepto",
            "Base Capital",
            "Tasa",
            "Interés",
            "Indexación/Sanciones",
            "Pago",
            "Saldo Capital",
            "Saldo Interés",
            "Saldo Total",
        ]
        if mostrar_saldo_a_favor:
            columnas_cronologia.append("Saldo a favor")
        tabla_cronologia = documento.add_table(rows=1, cols=len(columnas_cronologia))
        tabla_cronologia.style = "Table Grid"
        for celda, texto in zip(tabla_cronologia.rows[0].cells, columnas_cronologia, strict=True):
            _escribir_celda(
                celda, texto, color=self.c_cream, bold=True, font_name=_NOMBRE_FUENTE_EXTRABOLD
            )
            _sombrear_celda(celda, self.c_black_hex)
        for fila_datos in table_data:
            celdas_fila = tabla_cronologia.add_row().cells
            valores_fila = [
                fila_datos["fecha"],
                fila_datos["concepto"],
                fila_datos["base_capital"],
                fila_datos["tasa"],
                fila_datos["interes"],
                fila_datos["indexacion"],
                fila_datos["pago"],
                fila_datos["saldo_capital"],
                fila_datos["saldo_interes"],
                fila_datos["saldo_total"],
            ]
            if mostrar_saldo_a_favor:
                valores_fila.append(fila_datos.get("saldo_a_favor", "$0.00"))
            # Sprint 42: indicador visual (texto en rojo) para las filas cuya
            # obligacion de origen ya vencio su plazo de prescripcion/caducidad
            # (ReportTableBuilder.build_matrix ya expone "prescrita" por fila) --
            # no se excluyen de la tabla, solo se resaltan. Se usa un run
            # explicito (en vez de `celda.text = ...`) para poder colorear el
            # texto: el setter `.text` de python-docx no expone color de fuente.
            es_prescrita = bool(fila_datos.get("prescrita"))
            for celda, texto in zip(celdas_fila, valores_fila, strict=True):
                run = celda.paragraphs[0].add_run(texto)
                if es_prescrita:
                    run.font.color.rgb = self.c_prescrita
                    run.bold = True
                _sombrear_celda(celda, self.c_cream_hex)

        # Fijar ancho por columna (Sprint 50): "Table Grid" por defecto usa
        # autofit-to-contents, que Word resuelve al abrir el documento sin
        # respetar el ancho de pagina cuando el contenido es ancho (esa era la
        # causa real del desborde). `autofit = False` cambia la tabla a layout
        # fijo para que los anchos de abajo se respeten.
        tabla_cronologia.autofit = False
        anchos_cronologia = self._anchos_columnas_cronologia(
            columnas_cronologia, ancho_disponible
        )
        for indice, ancho in enumerate(anchos_cronologia):
            self._fijar_ancho_columna(tabla_cronologia, indice, ancho)

        if renta_liquida is not None:
            documento.add_paragraph()
            parrafo_renta_liquida = documento.add_paragraph()
            run_renta_liquida = parrafo_renta_liquida.add_run(
                "Depuración de Renta Líquida Gravable"
            )
            run_renta_liquida.bold = True
            run_renta_liquida.font.color.rgb = self.c_burgundy

            filas_renta_liquida = [
                ("Ingresos Netos", renta_liquida["ingresos_netos"]),
                ("Renta Bruta", renta_liquida["renta_bruta"]),
                ("Renta Líquida", renta_liquida["renta_liquida"]),
                ("¿Hubo Pérdida Líquida?", renta_liquida["hubo_perdida_liquida"]),
                ("RENTA LÍQUIDA GRAVABLE", renta_liquida["renta_liquida_gravable"]),
            ]
            tabla_renta_liquida = documento.add_table(rows=1, cols=2)
            tabla_renta_liquida.style = "Table Grid"
            celdas_encabezado_rl = tabla_renta_liquida.rows[0].cells
            for celda, texto in zip(celdas_encabezado_rl, ("Rubro", "Monto"), strict=True):
                _escribir_celda(
                    celda, texto, color=self.c_cream, bold=True, font_name=_NOMBRE_FUENTE_EXTRABOLD
                )
                _sombrear_celda(celda, self.c_black_hex)
            indice_ultima_fila_rl = len(filas_renta_liquida) - 1
            for indice, (etiqueta, valor) in enumerate(filas_renta_liquida):
                celdas_fila_rl = tabla_renta_liquida.add_row().cells
                es_fila_total = indice == indice_ultima_fila_rl
                for celda, texto in zip(celdas_fila_rl, (etiqueta, valor), strict=True):
                    _escribir_celda(
                        celda,
                        texto,
                        color=self.c_burgundy if es_fila_total else None,
                        bold=es_fila_total,
                    )
                    _sombrear_celda(celda, self.c_cream_hex)

        # Sprint 47: log de diferencias del recalculo historico post-Sprint-30
        # (memoriales de actualizacion/correccion y de correccion de error
        # aritmetico Art. 151 CPACA, ver app/engine/reports/memoriales.py) --
        # mismo patron aditivo que renta_liquida arriba (Sprint 15).
        if diferencia_recalculo is not None:
            documento.add_paragraph()
            parrafo_diferencia = documento.add_paragraph()
            run_diferencia = parrafo_diferencia.add_run(
                "Log de Diferencias — Recálculo Histórico (Sprint 30)"
            )
            run_diferencia.bold = True
            run_diferencia.font.color.rgb = self.c_burgundy

            filas_diferencia = [
                ("Liquidación Anterior", diferencia_recalculo["audit_log_anterior"]),
                (
                    "Valor Anterior (Sprint 30, pre-corrección)",
                    diferencia_recalculo["monto_anterior"],
                ),
                ("Valor Recalculado", diferencia_recalculo["monto_recalculado"]),
                ("Diferencia Recuperada", diferencia_recalculo["diferencia_monto"]),
                (
                    "Días Cubiertos (Anterior → Recalculado)",
                    f"{diferencia_recalculo['dias_cubiertos_anterior']} → "
                    f"{diferencia_recalculo['dias_cubiertos_recalculado']}",
                ),
            ]
            tabla_diferencia = documento.add_table(rows=1, cols=2)
            tabla_diferencia.style = "Table Grid"
            celdas_encabezado_diferencia = tabla_diferencia.rows[0].cells
            for celda, texto in zip(celdas_encabezado_diferencia, ("Rubro", "Valor"), strict=True):
                _escribir_celda(
                    celda, texto, color=self.c_cream, bold=True, font_name=_NOMBRE_FUENTE_EXTRABOLD
                )
                _sombrear_celda(celda, self.c_black_hex)
            indice_ultima_fila_diferencia = len(filas_diferencia) - 1
            for indice, (etiqueta, valor) in enumerate(filas_diferencia):
                celdas_fila_diferencia = tabla_diferencia.add_row().cells
                es_fila_total = indice == indice_ultima_fila_diferencia
                for celda, texto in zip(celdas_fila_diferencia, (etiqueta, valor), strict=True):
                    _escribir_celda(
                        celda,
                        texto,
                        color=self.c_burgundy if es_fila_total else None,
                        bold=es_fila_total,
                    )
                    _sombrear_celda(celda, self.c_cream_hex)

            documento.add_paragraph(diferencia_recalculo["resumen"])

        documento.save(self.output_path)
